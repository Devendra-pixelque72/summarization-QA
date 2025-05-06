"""
Async document processor for handling various document types.
"""
import os
import tempfile
from typing import Optional, Dict, Any, List, Tuple
import logging
import re
import asyncio
import aiofiles

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class AsyncDocumentProcessor:
    """Handle different document types for the summarizer with async support."""
    
    def __init__(self, verbose: bool = False):
        """
        Initialize the document processor.
        
        Args:
            verbose: Whether to log detailed information
        """
        self.verbose = verbose
        # Check for required libraries
        self._check_dependencies()
    
    def _check_dependencies(self) -> None:
        """Check if required libraries are installed."""
        missing_libs = []
        try:
            import PyPDF2
        except ImportError:
            missing_libs.append("PyPDF2")
            
        try:
            import docx
        except ImportError:
            missing_libs.append("python-docx")
            
        try:
            import markdown
        except ImportError:
            missing_libs.append("markdown")
            
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            missing_libs.append("beautifulsoup4")
            
        try:
            import aiofiles
        except ImportError:
            missing_libs.append("aiofiles")
        
        if missing_libs:
            missing_str = ", ".join(missing_libs)
            logger.warning(f"Missing libraries: {missing_str}")
            logger.warning("Some document types might not be processed correctly.")
            logger.warning(f"Install with: pip install {' '.join(missing_libs)}")
    
    async def process_document(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Process document based on file type and extract text asynchronously.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
            
        file_extension = os.path.splitext(file_path)[1].lower()
        
        # Process based on file extension
        if file_extension == '.pdf':
            return await self._process_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return await self._process_docx(file_path)
        elif file_extension == '.md':
            return await self._process_markdown(file_path)
        elif file_extension in ['.txt', '.text', '.csv', '.json', '.xml', '.html', '.htm']:
            return await self._process_text(file_path)
        else:
            logger.warning(f"Unsupported file type: {file_extension}, trying to process as text")
            return await self._process_text(file_path)
    
    async def _process_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process PDF files and extract text asynchronously."""
        try:
            import PyPDF2
            
            # PDF processing is CPU-bound, run in a thread pool
            def process_pdf():
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    num_pages = len(reader.pages)
                    
                    if self.verbose:
                        logger.info(f"Processing PDF with {num_pages} pages")
                    
                    # Extract text from all pages
                    text = ""
                    for page_num in range(num_pages):
                        page = reader.pages[page_num]
                        page_text = page.extract_text() or ""
                        text += page_text + "\n\n"
                    
                    # Remove excessive blank lines
                    text = re.sub(r'\n{3,}', '\n\n', text)
                    
                    # Get metadata
                    metadata = {
                        "file_type": "PDF",
                        "page_count": num_pages,
                        "word_count": len(text.split()),
                        "char_count": len(text)
                    }
                    
                    # Try to extract PDF metadata
                    if hasattr(reader, 'metadata') and reader.metadata:
                        pdf_info = reader.metadata
                        for key, value in pdf_info.items():
                            if isinstance(key, str) and isinstance(value, (str, int, float, bool)) and value:
                                clean_key = key.strip('/').lower()
                                metadata[f"pdf_{clean_key}"] = str(value)
                    
                    return text, metadata
            
            # Run the PDF processing in a thread pool
            loop = asyncio.get_event_loop()
            return await loop.run_in_executor(None, process_pdf)
                
        except Exception as e:
            logger.error(f"Error processing PDF: {str(e)}")
            # Fallback to basic text extraction
            return await self._process_text(file_path)
    
    async def _process_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process DOCX files and extract text asynchronously."""
        try:
            import docx
            
            # DOCX processing is CPU-bound, run in a thread pool
            def process_docx():
                doc = docx.Document(file_path)
                
                if self.verbose:
                    logger.info(f"Processing DOCX with {len(doc.paragraphs)} paragraphs")
                
                # Extract text
                full_text = []
                for para in doc.paragraphs:
                    if para.text.strip():  # Only add non-empty paragraphs
                        full_text.append(para.text)
                
                text = "\n\n".join(full_text)
                
                # Get metadata
                metadata = {
                    "file_type": "DOCX",
                    "paragraph_count": len(doc.paragraphs),
                    "word_count": len(text.split()),
                    "char_count": len(text),
                    "section_count": len(doc.sections)
                }
                
                # Try to get core properties
                try:
                    core_props = doc.core_properties
                    metadata["docx_title"] = core_props.title or "Unknown"
                    metadata["docx_author"] = core_props.author or "Unknown"
                    metadata["docx_created"] = str(core_props.created) if core_props.created else "Unknown"
                    metadata["docx_modified"] = str(core_props.modified) if core_props.modified else "Unknown"
                except Exception as e:
                    if self.verbose:
                        logger.warning(f"Could not extract DOCX properties: {str(e)}")
                
                return text, metadata
            
            # Run the DOCX processing in a thread pool
            loop = asyncio.get_event_loop()
            return await loop.run_in_executor(None, process_docx)
            
        except Exception as e:
            logger.error(f"Error processing DOCX: {str(e)}")
            # Fallback to basic text extraction
            return await self._process_text(file_path)
    
    async def _process_markdown(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process Markdown files and extract text asynchronously."""
        try:
            import markdown
            from bs4 import BeautifulSoup
            
            try:
                import aiofiles
                async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                    md_content = await file.read()
            except ImportError:
                # Fallback to synchronous reading if aiofiles is not available
                with open(file_path, 'r', encoding='utf-8') as file:
                    md_content = file.read()
            
            if self.verbose:
                logger.info(f"Processing Markdown file with {len(md_content.splitlines())} lines")
            
            # Convert markdown to HTML (CPU-bound, run in thread pool)
            def process_markdown():
                # Convert markdown to HTML
                html = markdown.markdown(md_content)
                
                # Use BeautifulSoup for better HTML to text conversion
                soup = BeautifulSoup(html, features="html.parser")
                text = soup.get_text("\n\n")
                
                # Clean up extra whitespace
                text = re.sub(r'\n{3,}', '\n\n', text)
                text = re.sub(r' {2,}', ' ', text)
                
                # Get metadata
                line_count = len(md_content.splitlines())
                word_count = len(text.split())
                
                metadata = {
                    "file_type": "Markdown",
                    "line_count": line_count,
                    "word_count": word_count,
                    "char_count": len(text)
                }
                
                return text, metadata
            
            # Run the Markdown processing in a thread pool
            loop = asyncio.get_event_loop()
            return await loop.run_in_executor(None, process_markdown)
            
        except Exception as e:
            logger.error(f"Error processing Markdown: {str(e)}")
            # Fallback to basic text extraction
            return await self._process_text(file_path)
    
    async def _process_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process plain text files asynchronously."""
        try:
            try:
                import aiofiles
                async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                    text = await file.read()
            except ImportError:
                # Fallback to synchronous reading if aiofiles is not available
                with open(file_path, 'r', encoding='utf-8') as file:
                    text = file.read()
            
            if self.verbose:
                logger.info(f"Processing text file with {len(text.splitlines())} lines")
            
            # Get metadata
            line_count = len(text.splitlines())
            word_count = len(text.split())
            
            metadata = {
                "file_type": "Text",
                "line_count": line_count,
                "word_count": word_count,
                "char_count": len(text)
            }
            
            return text, metadata
            
        except UnicodeDecodeError:
            # Try different encodings if utf-8 fails
            try:
                try:
                    import aiofiles
                    async with aiofiles.open(file_path, 'r', encoding='latin-1') as file:
                        text = await file.read()
                except ImportError:
                    # Fallback to synchronous reading
                    with open(file_path, 'r', encoding='latin-1') as file:
                        text = file.read()
                
                if self.verbose:
                    logger.info(f"Processed text file with latin-1 encoding")
                
                line_count = len(text.splitlines())
                word_count = len(text.split())
                
                metadata = {
                    "file_type": "Text",
                    "line_count": line_count,
                    "word_count": word_count,
                    "char_count": len(text),
                    "encoding": "latin-1"
                }
                
                return text, metadata
            except Exception as e:
                logger.error(f"Error processing text file with latin-1 encoding: {str(e)}")
                raise ValueError(f"Could not read file with any encoding: {file_path}")
        except Exception as e:
            logger.error(f"Error processing text file: {str(e)}")
            raise ValueError(f"Error processing text file: {str(e)}")